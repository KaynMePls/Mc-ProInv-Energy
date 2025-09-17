from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify
from app.models import TransaccionContable, Compra, Usuario, Producto
from flask_login import login_required, current_user
from sqlalchemy import func, extract
from datetime import datetime
from dateutil.relativedelta import relativedelta
from app import db
from app.forms import TransaccionForm, CompraForm, EstadoCompraForm
import pandas as pd
from io import BytesIO
from flask import send_file
import pdfkit
from flask_wtf.csrf import CSRFProtect, CSRFError, validate_csrf
from docx import Document
from docx.shared import Pt
from decimal import Decimal
from docx.shared import Pt, RGBColor
from flask_wtf import FlaskForm
from flask import send_file, request, url_for


contabilidad_bp = Blueprint('contabilidad', __name__, url_prefix='/contabilidad')

@contabilidad_bp.route('/dashboard')
@login_required
def dashboard():
    # Obtener filtros del formulario
    mes = request.args.get('mes')
    anio = request.args.get('anio')
    categoria = request.args.get('categoria')
    buscar = request.args.get('buscar')

    # Base query
    query = TransaccionContable.query

    # Aplicar filtros
    if mes:
        query = query.filter(extract('month', TransaccionContable.fecha) == int(mes))
    if anio:
        query = query.filter(extract('year', TransaccionContable.fecha) == int(anio))
    if categoria:
        query = query.filter(TransaccionContable.categoria == categoria)
    if buscar:
        query = query.filter(TransaccionContable.descripcion.ilike(f'%{buscar}%'))

    # Resumen financiero filtrado
    total_ingresos = query.filter_by(tipo='ingreso').with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
    total_egresos = query.filter_by(tipo='egreso').with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
    balance_neto = (total_ingresos or 0) - (total_egresos or 0)

    # Últimas transacciones filtradas
    ultimas_transacciones = query.order_by(TransaccionContable.fecha.desc()).limit(10).all()

    # Gráfico de evolución mensual (últimos 6 meses)
    labels_meses = []
    datos_ingresos = []
    datos_egresos = []

    for i in range(5, -1, -1):
        mes_dt = datetime.now().replace(day=1) - relativedelta(months=i)
        mes_label = mes_dt.strftime('%b %Y')
        labels_meses.append(mes_label)
        ingresos_mes = query.filter(
            TransaccionContable.tipo == 'ingreso',
            extract('month', TransaccionContable.fecha) == mes_dt.month,
            extract('year', TransaccionContable.fecha) == mes_dt.year
        ).with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
        egresos_mes = query.filter(
            TransaccionContable.tipo == 'egreso',
            extract('month', TransaccionContable.fecha) == mes_dt.month,
            extract('year', TransaccionContable.fecha) == mes_dt.year
        ).with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
        datos_ingresos.append(float(ingresos_mes or 0))
        datos_egresos.append(float(egresos_mes or 0))

    # Gráfico de egresos por categoría (pie)
    categorias_db = [c[0] for c in query.with_entities(TransaccionContable.categoria).filter_by(tipo='egreso').distinct()]
    labels_egresos = categorias_db
    datos_egresos_categoria = [
        float(query.filter_by(tipo='egreso', categoria=cat).with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar() or 0)
        for cat in categorias_db
    ]

    # Para el filtro de categorías en el select
    todas_categorias = [c[0] for c in TransaccionContable.query.with_entities(TransaccionContable.categoria).distinct()]

    return render_template(
        'empleado/contabilidad/contabilidad.dashboard.html',
        total_ingresos=total_ingresos,
        total_egresos=total_egresos,
        balance_neto=balance_neto,
        ultimas_transacciones=ultimas_transacciones,
        labels_meses=labels_meses,
        datos_ingresos=datos_ingresos,
        datos_egresos=datos_egresos,
        labels_egresos=labels_egresos,
        datos_egresos_categoria=datos_egresos_categoria,
        categorias=todas_categorias,
        now=datetime.now()
    )

# Ruta AJAX para categorías dinámicas
@contabilidad_bp.route('/categorias_por_tipo')
@login_required
def categorias_por_tipo():
    tipo = request.args.get('tipo')
    if tipo == 'ingreso':
        categorias = [
            {"id": "Ventas", "nombre": "Ventas"},
            {"id": "Intereses", "nombre": "Intereses"},
            {"id": "Devoluciones", "nombre": "Devoluciones"},
            {"id": "Inversiones", "nombre": "Inversiones"},
            {"id": "Otros ingresos", "nombre": "Otros ingresos"}
        ]
    else:
        categorias = [
            {"id": "Compras", "nombre": "Compras"},
            {"id": "Servicios", "nombre": "Servicios"},
            {"id": "Nómina", "nombre": "Nómina"},
            {"id": "Impuestos", "nombre": "Impuestos"},
            {"id": "Mantenimiento", "nombre": "Mantenimiento"},
            {"id": "Otros egresos", "nombre": "Otros egresos"}
        ]
    return jsonify(categorias)

# Ruta para la nueva transacción contable
@contabilidad_bp.route('/nueva_transaccion', methods=['GET', 'POST'])
@login_required
def nueva_transaccion():
    form = TransaccionForm()
    cancel_url = request.args.get('cancel_url') or url_for('contabilidad.dashboard')

    # Asegura opciones de categoría según el tipo (por defecto 'ingreso')
    tipo_actual = form.tipo.data or 'ingreso'
    if tipo_actual == 'ingreso':
        form.categoria.choices = [
            ("Ventas", "Ventas"),
            ("Intereses", "Intereses"),
            ("Devoluciones", "Devoluciones"),
            ("Inversiones", "Inversiones"),
            ("Otros ingresos", "Otros ingresos")
        ]
    else:
        form.categoria.choices = [
            ("Compras", "Compras"),
            ("Servicios", "Servicios"),
            ("Nómina", "Nómina"),
            ("Impuestos", "Impuestos"),
            ("Mantenimiento", "Mantenimiento"),
            ("Otros egresos", "Otros egresos")
        ]

    if form.validate_on_submit():
        transaccion = TransaccionContable(
            tipo=form.tipo.data,
            monto=form.monto.data,
            descripcion=form.descripcion.data,
            categoria=form.categoria.data,
            fecha=form.fecha.data,
            usuario_id=current_user.id
        )
        db.session.add(transaccion)
        db.session.commit()
        flash('Transacción registrada exitosamente.', 'success')
        return redirect(cancel_url)

    return render_template(
        'empleado/contabilidad/contabilidad.nueva_transaccion.html',
        form=form,
        cancel_url=cancel_url
    )

# Ruta para transacciones
@contabilidad_bp.route('/transacciones')
@login_required
def transacciones():
    # Filtros
    tipo = request.args.get('tipo')
    categoria = request.args.get('categoria')
    desde = request.args.get('desde')
    hasta = request.args.get('hasta')
    page = request.args.get('page', 1, type=int)
    per_page = 10

    query = TransaccionContable.query

    if tipo:
        query = query.filter_by(tipo=tipo)
    if categoria:
        query = query.filter(TransaccionContable.categoria.ilike(f"%{categoria}%"))
    if desde:
        query = query.filter(TransaccionContable.fecha >= desde)
    if hasta:
        query = query.filter(TransaccionContable.fecha <= hasta)

    pagination = query.order_by(TransaccionContable.fecha.desc()).paginate(page=page, per_page=per_page, error_out=False)
    transacciones = pagination.items

    # Totales filtrados
    total_ingresos = query.filter_by(tipo='ingreso').with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
    total_egresos = query.filter_by(tipo='egreso').with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
    balance_neto = (total_ingresos or 0) - (total_egresos or 0)

    # Gráfico de evolución mensual (últimos 6 meses)
    labels_meses = []
    datos_ingresos = []
    datos_egresos = []

    for i in range(5, -1, -1):
        mes = datetime.now().replace(day=1) - relativedelta(months=i)
        mes_label = mes.strftime('%b %Y')
        labels_meses.append(mes_label)
        ingresos_mes = query.filter(
            TransaccionContable.tipo == 'ingreso',
            extract('month', TransaccionContable.fecha) == mes.month,
            extract('year', TransaccionContable.fecha) == mes.year
        ).with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
        egresos_mes = query.filter(
            TransaccionContable.tipo == 'egreso',
            extract('month', TransaccionContable.fecha) == mes.month,
            extract('year', TransaccionContable.fecha) == mes.year
        ).with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
        datos_ingresos.append(float(ingresos_mes or 0))
        datos_egresos.append(float(egresos_mes or 0))

    return render_template(
        'empleado/contabilidad/contabilidad.transacciones.html',
        transacciones=transacciones,
        total_ingresos=total_ingresos,
        total_egresos=total_egresos,
        balance_neto=balance_neto,
        page=page,
        has_next=pagination.has_next,
        labels_meses=labels_meses,
        datos_ingresos=datos_ingresos,
        datos_egresos=datos_egresos
    )

# Ruta para reportes de contabilidad
@contabilidad_bp.route('/reporte')
@login_required
def reporte():
    transacciones = TransaccionContable.query.order_by(TransaccionContable.fecha.desc()).all()
    total_ingresos = TransaccionContable.query.filter_by(tipo='ingreso').with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
    total_egresos = TransaccionContable.query.filter_by(tipo='egreso').with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
    balance_neto = total_ingresos - total_egresos
    porcentaje_egreso = Decimal('0.7')  # <-- Agrega esta línea

    return render_template(
        'empleado/contabilidad/contabilidad.reporte.html',
        transacciones=transacciones,
        total_ingresos=total_ingresos,
        total_egresos=total_egresos,
        balance_neto=balance_neto,
        porcentaje_egreso=porcentaje_egreso,  # <-- Pásala al template
        now=datetime.now()
    )

# Ruta para exportar transacciones a Excel
@contabilidad_bp.route('/exportar_excel')
@login_required
def exportar_excel():
    tipo = request.args.get('tipo')
    categoria = request.args.get('categoria')
    desde = request.args.get('desde')
    hasta = request.args.get('hasta')

    query = TransaccionContable.query

    if tipo:
        query = query.filter_by(tipo=tipo)
    if categoria:
        query = query.filter(TransaccionContable.categoria.ilike(f"%{categoria}%"))
    if desde:
        query = query.filter(TransaccionContable.fecha >= desde)
    if hasta:
        query = query.filter(TransaccionContable.fecha <= hasta)

    transacciones = query.order_by(TransaccionContable.fecha.desc()).all()

    # Prepara los datos para Excel
    data = [{
        'Fecha': t.fecha.strftime('%Y-%m-%d'),
        'Tipo': t.tipo.capitalize(),
        'Monto': t.monto,
        'Categoría': t.categoria,
        'Descripción': t.descripcion
    } for t in transacciones]

    df = pd.DataFrame(data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Transacciones')
        workbook  = writer.book
        worksheet = writer.sheets['Transacciones']

        # Formatos
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#0d6efd',
            'font_color': 'white',
            'border': 1,
            'align': 'center'
        })
        money_format = workbook.add_format({'num_format': '$#,##0.00', 'border': 1})
        text_format = workbook.add_format({'border': 1})
        tipo_format = workbook.add_format({'border': 1, 'bg_color': '#d1e7dd'})
        egreso_format = workbook.add_format({'border': 1, 'bg_color': '#f8d7da'})

        # Encabezados en negrita y color
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            worksheet.set_column(col_num, col_num, 18)

        # Formato de filas
        for row_num, t in enumerate(transacciones, start=1):
            worksheet.write(row_num, 0, t.fecha.strftime('%Y-%m-%d'), text_format)
            # Tipo con color
            if t.tipo == 'ingreso':
                worksheet.write(row_num, 1, t.tipo.capitalize(), tipo_format)
            else:
                worksheet.write(row_num, 1, t.tipo.capitalize(), egreso_format)
            worksheet.write(row_num, 2, t.monto, money_format)
            worksheet.write(row_num, 3, t.categoria, text_format)
            worksheet.write(row_num, 4, t.descripcion, text_format)

        # Autofiltro
        worksheet.autofilter(0, 0, len(df), len(df.columns) - 1)

    output.seek(0)
    return send_file(
        output,
        download_name='transacciones.xlsx',
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


# Ruta para editar una transacción
@contabilidad_bp.route('/editar_transaccion/<int:transaccion_id>', methods=['GET', 'POST'])
@login_required
def editar_transaccion(transaccion_id):
    transaccion = TransaccionContable.query.get_or_404(transaccion_id)
    form = TransaccionForm(obj=transaccion)

    # Opciones dinámicas para categoría según tipo
    if form.tipo.data == 'ingreso':
        form.categoria.choices = [
            ("Ventas", "Ventas"),
            ("Intereses", "Intereses"),
            ("Devoluciones", "Devoluciones"),
            ("Inversiones", "Inversiones"),
            ("Otros ingresos", "Otros ingresos")
        ]
    else:
        form.categoria.choices = [
            ("Compras", "Compras"),
            ("Servicios", "Servicios"),
            ("Nómina", "Nómina"),
            ("Impuestos", "Impuestos"),
            ("Mantenimiento", "Mantenimiento"),
            ("Otros egresos", "Otros egresos")
        ]

    if form.validate_on_submit():
        transaccion.tipo = form.tipo.data
        transaccion.monto = form.monto.data
        transaccion.fecha = form.fecha.data
        transaccion.categoria = form.categoria.data
        transaccion.descripcion = form.descripcion.data
        db.session.commit()
        flash('Transacción actualizada exitosamente.', 'success')
        return redirect(url_for('contabilidad.transacciones'))

    return render_template(
        'empleado/contabilidad/contabilidad.nueva_transaccion.html',
        form=form,
        cancel_url=request.args.get('cancel_url') or url_for('contabilidad.dashboard')
    )

# Ruta para eliminar una transacción
@contabilidad_bp.route('/eliminar_transaccion/<int:transaccion_id>', methods=['POST'])
@login_required
def eliminar_transaccion(transaccion_id):
    try:
        validate_csrf(request.form.get('csrf_token'))
    except CSRFError:
        flash('Token CSRF inválido o ausente.', 'danger')
        return redirect(url_for('contabilidad.transacciones'))

    transaccion = TransaccionContable.query.get_or_404(transaccion_id)
    db.session.delete(transaccion)
    db.session.commit()
    flash('Transacción eliminada exitosamente.', 'success')
    return redirect(url_for('contabilidad.transacciones'))


# Ruta para exportar transacciones a Word
@contabilidad_bp.route('/exportar_word')
@login_required
def exportar_word():
    transacciones = TransaccionContable.query.order_by(TransaccionContable.fecha.desc()).all()
    total_ingresos = TransaccionContable.query.filter_by(tipo='ingreso').with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
    total_egresos = TransaccionContable.query.filter_by(tipo='egreso').with_entities(func.coalesce(func.sum(TransaccionContable.monto), 0)).scalar()
    balance_neto = total_ingresos - total_egresos
    porcentaje_egreso = Decimal('0.7')
    now = datetime.now()

    doc = Document()
    doc.add_heading('Reporte Contable', 0)

    # Resumen financiero con formato
    resumen = doc.add_table(rows=4, cols=2)
    resumen.style = 'Light Shading Accent 1'
    resumen.autofit = True
    resumen.cell(0, 0).text = 'Generado el:'
    resumen.cell(0, 1).text = now.strftime("%d/%m/%Y %H:%M")
    resumen.cell(1, 0).text = 'Ingresos'
    resumen.cell(1, 1).text = f"${total_ingresos:,.0f}"
    resumen.cell(2, 0).text = 'Egresos'
    resumen.cell(2, 1).text = f"${total_egresos:,.0f}"
    resumen.cell(3, 0).text = 'Balance Neto'
    resumen.cell(3, 1).text = f"${balance_neto:,.0f}"

    for i in range(4):
        for j in range(2):
            resumen.cell(i, j).paragraphs[0].runs[0].font.size = Pt(12)
            resumen.cell(i, j).paragraphs[0].runs[0].font.bold = True if j == 0 else False

    doc.add_paragraph()  # Espacio

    doc.add_heading('Transacciones', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fecha'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Monto'
    hdr_cells[3].text = 'Categoría'
    hdr_cells[4].text = 'Descripción'

    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)

    for t in transacciones:
        row_cells = table.add_row().cells
        row_cells[0].text = t.fecha.strftime('%Y-%m-%d')
        row_cells[1].text = t.tipo.capitalize()
        row_cells[2].text = f"${t.monto:,.0f}"
        row_cells[3].text = t.categoria
        row_cells[4].text = t.descripcion[:60] + ("..." if len(t.descripcion) > 60 else "")

        # Formato condicional para tipo
        tipo_color = '00B050' if t.tipo == 'ingreso' else 'C00000'
        for run in row_cells[1].paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(tipo_color)
            run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph("Este reporte fue generado automáticamente por el sistema contable Solarium. Para soporte, escriba a info@solarium.com.")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(output, download_name='reporte_contable.docx', as_attachment=True)

# Ruta para registrar una compra a proveedores
@contabilidad_bp.route('/nueva_compra', methods=['GET', 'POST'])
@login_required
def nueva_compra():
    form = CompraForm()
    proveedores = Usuario.query.filter_by(rol='proveedor').all()
    proveedores_productos = {
        p.id: Producto.query.filter_by(proveedor_id=p.id).all()
        for p in proveedores
    }
    usuarios_dict = {p.id: p for p in proveedores}
    form.proveedor.choices = [(p.nombre, p.nombre) for p in proveedores]

    cancel_url = request.args.get('cancel_url') or url_for('contabilidad.dashboard')

    if form.validate_on_submit():
        compra = Compra(
            proveedor=form.proveedor.data,
            producto=form.producto.data,
            cantidad=form.cantidad.data,
            precio_unitario=form.precio_unitario.data,
            fecha=form.fecha.data,
            usuario_id=current_user.id
        )
        db.session.add(compra)
        db.session.commit()
        flash('Compra registrada exitosamente.', 'success')
        return redirect(cancel_url)
    return render_template(
        'empleado/contabilidad/contabilidad.nueva_compra.html',
        form=form,
        proveedores_productos=proveedores_productos,
        usuarios_dict=usuarios_dict,
        Decimal=Decimal,
        cancel_url=cancel_url
    )
# Ruta para Comprar productos
@contabilidad_bp.route('/comprar_productos', methods=['GET', 'POST'])
@login_required
def comprar_productos():
    proveedores = Usuario.query.filter_by(rol='proveedor').all()
    proveedores_productos = {p.id: Producto.query.filter_by(proveedor_id=p.id).all() for p in proveedores}
    usuarios_dict = {p.id: p for p in proveedores}
    form = CompraForm()
    return render_template(
        'empleado/contabilidad/contabilidad.nueva_compra.html',
        proveedores_productos=proveedores_productos,
        usuarios_dict=usuarios_dict,
        form=form,
        Decimal=Decimal
    )
# Ruta para generar ordenes de compra

@contabilidad_bp.route('/generar_orden_compra/<int:producto_id>', methods=['POST'])
@login_required
def generar_orden_compra(producto_id):
    producto = Producto.query.get_or_404(producto_id)
    cantidad = int(request.form.get('cantidad', 1))
    precio_venta = producto.precio * Decimal('0.7')  # Aplica el 30% de descuento
    proveedor = Usuario.query.get(producto.proveedor_id)  # Busca el usuario proveedor

    compra = Compra(
        proveedor=proveedor,              # Objeto Usuario, NO string
        producto=producto.nombre,         # Nombre del producto (string)
        cantidad=cantidad,
        precio_unitario=precio_venta,
        fecha=datetime.now(),
        usuario_id=current_user.id
    )
    db.session.add(compra)
    # Actualiza el stock del producto
    producto.stock = (producto.stock or 0) + cantidad
    db.session.commit()
    flash('Orden de compra generada y stock actualizado.', 'success')
    return redirect(url_for('contabilidad.nueva_compra'))
# Ruta para ver compras
@contabilidad_bp.route('/compras')
@login_required
def compras():
    compras = Compra.query.order_by(Compra.fecha.desc()).all()
    form = EstadoCompraForm()
    return render_template('empleado/contabilidad/contabilidad.compras.html', compras=compras, form=form)

# Ruta estado de compra
@contabilidad_bp.route('/cambiar_estado_compra/<int:compra_id>', methods=['POST'])
@login_required
def cambiar_estado_compra(compra_id):
    compra = Compra.query.get_or_404(compra_id)
    nuevo_estado = request.form.get('estado')
    if nuevo_estado in ['aprobada', 'rechazada']:
        proveedor = compra.proveedor  # Ya es el objeto Usuario
        producto = Producto.query.filter_by(nombre=compra.producto, proveedor_id=proveedor.id if proveedor else None).first()
        if producto and compra.precio_unitario >= producto.precio:
            flash('El precio de compra no puede ser mayor o igual al precio de venta.', 'danger')
            return redirect(url_for('contabilidad.compras'))

        compra.estado = nuevo_estado
        db.session.commit()
        flash(f'Orden de compra marcada como {nuevo_estado}.', 'success')
        # Si se aprueba, crea la transacción contable y publica el producto
        if nuevo_estado == 'aprobada':
            transaccion = TransaccionContable(
                tipo='egreso',
                monto=compra.precio_unitario * compra.cantidad,
                descripcion=f'Compra aprobada: {compra.producto} a {proveedor.nombre if proveedor else ""}',
                categoria='compras',
                usuario_id=compra.usuario_id
            )
            db.session.add(transaccion)
            if producto:
                producto.publicado = True
            db.session.commit()
    else:
        flash('Estado inválido.', 'danger')
    return redirect(url_for('contabilidad.compras'))